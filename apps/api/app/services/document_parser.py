"""
Document Parser Service
Extracts text from PDF and DOCX files
"""
import io
import logging
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# OCR imports (optional - for image-based PDFs)
import os

try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_bytes

    # Set Tesseract executable path (Windows)
    # Try common installation locations
    tesseract_paths = [
        r"C:\Users\2000147382\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Programs', 'Tesseract-OCR', 'tesseract.exe'),
    ]

    # Find the first valid path
    tesseract_found = False
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            print(f"✅ Tesseract found at: {path}")
            logger.info(f"✅ Tesseract found at: {path}")
            tesseract_found = True
            break

    if not tesseract_found:
        print("⚠️ Tesseract not found in common locations. OCR may not work.")
        logger.warning("⚠️ Tesseract not found in common locations")

    OCR_AVAILABLE = True
    print("✅ OCR libraries available - can process image-based PDFs")
    logger.info("✅ OCR libraries available - can process image-based PDFs")
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"⚠️ OCR libraries not available: {e}")
    print("   Install: pip install pytesseract pdf2image Pillow")
    logger.warning(f"⚠️ OCR libraries not available: {e}")


class DocumentParser:
    """Service to extract text from various document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file using multiple methods for better accuracy
        """
        text = ""
        
        # Method 1: Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info("Extracted text using pdfplumber")
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 2: Fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                logger.info("Extracted text using PyPDF2")
                return text.strip()
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
        
        # Method 3: OCR for image-based PDFs (if available)
        if not text.strip() and OCR_AVAILABLE:
            try:
                logger.info("PDF appears to be image-based. Attempting OCR...")

                # Set poppler path for Windows (if needed)
                poppler_path = None
                possible_poppler_paths = [
                    r"C:\poppler\Library\bin",
                    r"C:\Program Files\poppler\Library\bin",
                    os.path.join(os.environ.get('PROGRAMFILES', ''), 'poppler', 'Library', 'bin'),
                ]

                for path in possible_poppler_paths:
                    if os.path.exists(path):
                        poppler_path = path
                        logger.info(f"✅ Poppler found at: {path}")
                        break

                # Convert PDF to images
                if poppler_path:
                    images = convert_from_bytes(file_content, poppler_path=poppler_path)
                else:
                    # Try without explicit path (if in system PATH)
                    images = convert_from_bytes(file_content)

                ocr_text = ""
                for i, image in enumerate(images):
                    logger.info(f"  OCR processing page {i+1}/{len(images)}...")
                    page_text = pytesseract.image_to_string(image)
                    if page_text:
                        ocr_text += page_text + "\n"

                if ocr_text.strip():
                    logger.info(f"✅ OCR extracted {len(ocr_text)} characters from {len(images)} pages")
                    return ocr_text.strip()
                else:
                    logger.warning("OCR completed but extracted no text")

            except Exception as e:
                logger.error(f"OCR extraction failed: {e}")

        # If all methods fail
        if not text.strip():
            if OCR_AVAILABLE:
                logger.error("❌ Failed to extract text from PDF (tried text extraction + OCR)")
            else:
                logger.error("❌ Failed to extract text from PDF. This might be an image-based PDF. Install OCR: pip install pytesseract pdf2image Pillow")
            return ""

        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            logger.info("Extracted text from DOCX")
            return text.strip()
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    @staticmethod
    def extract_text(file_content: bytes, content_type: str) -> Optional[str]:
        """
        Extract text from document based on content type
        
        Args:
            file_content: File bytes
            content_type: MIME type of the file
        
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            if content_type == "application/pdf":
                return DocumentParser.extract_text_from_pdf(file_content)
            
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                return DocumentParser.extract_text_from_docx(file_content)
            
            else:
                logger.warning(f"Unsupported content type: {content_type}")
                return None
        
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return None


# Global instance
document_parser = DocumentParser()
